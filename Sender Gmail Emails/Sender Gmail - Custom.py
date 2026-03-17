import os,sys  #Standard Python Libraries
import pandas as pd
import docx 
from docxtpl import DocxTemplate, InlineImage  # pip install docxtpl
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from pathlib import Path
import comtypes.client
import mammoth
import html as html_lib



current_dir = os.getcwd() + '\Data\Asis_curso - Copy.xlsx'
context = pd.read_excel(current_dir)
context['Nombre'] = context['Nombre'].str.replace("-", " ").str.upper()
context['Apellido'] = context['Apellido'].str.replace("-", " ").str.upper()


number_of_rows=context.shape[0]


# Setup port number and server name
smtp_port = 587              # Standard secure SMTP port
smtp_server = "smtp.gmail.com"  # Google SMTP Server



# Set up the email lists
email_from = 'edgaus98@gmail.com'
# Define the password (better to reference externally)
pswd = 'ywqbodbjvxvmxtqy' 

# Define the email function (dont call it email!)


def sender_html(person, filename,i,pdf_name,mails,INDEXAS):

    Name = person['Nombre']
    cc = 'edgar.agustin@cinvestav.mx'


    for mail in mails:
        
        # make a MIME object to define parts of the email
        msg = MIMEMultipart("alternatives")
        msg['From'] = 'Capítulo Estudiantil'
        msg['To'] = mail
        msg['Cc'] = cc
        msg['Subject'] = "Certificado Curso Fotolumiscenia en Nanomateriales"
 

        with open(INDEXAS, 'r', encoding="utf-8") as file:
            html_text = file.read()

# 2. GUARDAR (Opcional para debug): También con utf-8
        with open('html_text.html', 'w', encoding="utf-8") as file:
            file.write(html_text)

# 3. ADJUNTAR: Especifica explícitamente el charset en MIMEText
# Esto le dice al servidor de correo (Gmail, Outlook) cómo interpretar los acentos
        htmlPart = MIMEText(html_text, 'html', 'utf-8') 
        msg.attach(htmlPart)

        attachment= open(filename, 'rb')  # r for read and b for binary

    # Encode as base 64
        attachment_package = MIMEBase('application', 'octet-stream')
        attachment_package.set_payload((attachment).read())
        encoders.encode_base64(attachment_package)
        attachment_package.add_header('Content-Disposition', "attachment; filename= " + pdf_name)
        msg.attach(attachment_package)
    
    # Cast as string
        text = msg.as_string()

    # Connect with the server
        print("Connecting to server...")
        TIE_server = smtplib.SMTP(smtp_server, smtp_port)
        TIE_server.starttls()
        TIE_server.login(email_from,pswd)
        print("Succesfully connected to server")
        print()


    # Send emails to "person" as list is iterated
        print(f"Sending email to: {mail}...")
        TIE_server.sendmail(email_from, [mail, cc], text)
        print(f"Email sent to: {Name}")
        print()


        print(f'Correo número: {i+1}')
        print()

    # Close the port
        TIE_server.quit()

def extraer_subcadena(cadena, caracter_inicio, caracter_fin):
    subcadenas = []
    while caracter_inicio in cadena and caracter_fin in cadena:
        inicio = cadena.index(caracter_inicio) + 1
        fin = cadena.index(caracter_fin, inicio)
        subcadenas.append(cadena[inicio:fin])
        cadena = cadena[fin + 1:]
    return subcadenas

limit = context.shape[0]   

for i in range(142,limit):
    
    
    extract = context.loc[i]

    if extract['CONSTANCIA (1 SI)'] ==1:

        authors_email=[ extract['Correo electrónico']  ]
        
        version = 'Data\Template Foto Curso.docx'
        doc = DocxTemplate(version)
        doc.render(extract)

        # Set the paths for the Word and PDF files
        word_path = f'Data\Documents_Generated\Word\Acceptance Letter-{i}.docx'
        pdf_path = f'Data\Documents_Generated\PDF\Acceptance Letter-{i}.pdf'
        pdf_name = f"Certificate Id-{i}.pdf"


        doc.save(word_path)

        # Set the paths for the Word and PDF files
        word_path = f'Data\Documents_Generated\Word\Acceptance Letter-{i}.docx'
        pdf_path = f'Data\Documents_Generated\PDF\Acceptance Letter-{i}.pdf'
        pdf_name = f"Certificado Id-{i}.pdf"
 
        # Load the Word document using the docx library
        doc = docx.Document(word_path)
 
        # Save the Word document as a PDF using Microsoft Word
        word = comtypes.client.CreateObject("Word.Application")
        docx_path = os.path.abspath(word_path)
        pdf_path = os.path.abspath(pdf_path)

 
        pdf_format = 17  # PDF file format code
        word.Visible = False
        in_file = word.Documents.Open(docx_path)
        in_file.SaveAs(pdf_path, FileFormat=pdf_format)
        in_file.Close()
  
        # Quit Microsoft Word
        word.Quit()


# 1. Usa rutas RAW (r'') para evitar errores de caracteres en Windows
        version = r'Data\html.docx'
        version_SAVE = r'Data\html_SAVE.docx'

# 2. Renderizar la plantilla
        doc_tpl = DocxTemplate(version)
        doc_tpl.render(extract)
        doc_tpl.save(version_SAVE)

# 3. Extraer el texto directamente (Sin pasar por un .txt intermedio)
# Esto garantiza que los acentos se mantengan en la memoria de Python
        doc_result = docx.Document(version_SAVE)
        full_text = []

        for p in doc_result.paragraphs:
            full_text.append(p.text)

# Unimos todo el texto con saltos de línea
        text = "\n".join(full_text)

        nombre_archivo = "resultado_texto.txt"

# Guardamos con el encoding correcto
        with open(nombre_archivo, "w", encoding="utf-8") as f:
            f.write(text)

        # Quit Microsoft Word
        
    
        # Run the function
        sender_html(extract,pdf_path,i,pdf_name,authors_email,'resultado_texto.txt')
    
        
        
