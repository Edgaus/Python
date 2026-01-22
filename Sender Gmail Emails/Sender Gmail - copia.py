import os,sys  #Standard Python Libraries
import pandas as pd
import docx 
from docxtpl import DocxTemplate, InlineImage  # pip install docxtpl
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from pathlib import Path
import comtypes.client
import mammoth
import html as html_lib



current_dir = os.getcwd() + '\Datos\Sender Gmail Emails\Practica.xlsx'
context = pd.read_excel(current_dir)
context['Title'] = context['Title'].str.upper()
context['Authors'] = context['Authors'].str.replace("-", " ").str.upper()
context['autor_todos'] = context['autor_todos'].str.replace("-", " ").str.upper()


number_of_rows=context.shape[0]


# Setup port number and server name
smtp_port = 587              # Standard secure SMTP port
smtp_server = "smtp.gmail.com"  # Google SMTP Server



# Set up the email lists
email_from = 'edgaus98@gmail.com'
# Define the password (better to reference externally)
pswd = 'asdasdasd' 

# Define the email function (dont call it email!)


def sender_html(person, filename,i,pdf_name,mails,INDEXAS):

    Name = person['Authors']

    for mail in mails:

        # make a MIME object to define parts of the email
        msg = MIMEMultipart("alternatives")
        msg['From'] = 'Sociedad Mexicana de Ciencia y Tencologia de Superficies y Materiales A.C.'
        msg['To'] = 'edgaus98@gmail.com'
        msg['Subject'] = "Certificate of contribution at the XVII-International Conference on Surfaces Materials and Vacuum"
 

    # Attach the body of the message
        with open(INDEXAS, 'r') as file:
            html_text = file.read()

        with open('html_text.html', 'w', encoding="utf-8") as file:
            file.write(html_text)
        htmlPart = MIMEText(html_text, 'html' )
        msg.attach(htmlPart)

    
    # Open the file in python as a binary
        attachment= open(filename, 'rb')  # r for read and b for binary

    # Encode as base 64
        attachment_package = MIMEBase('application', 'octet-stream')
        attachment_package.set_payload((attachment).read())
        encoders.encode_base64(attachment_package)
        attachment_package.add_header('Content-Disposition', "attachment; filename= " + pdf_name)
        msg.attach(attachment_package)
    
    # Cast as string
        text = msg.as_string()

    # Connect with the server
        print("Connecting to server...")
        TIE_server = smtplib.SMTP(smtp_server, smtp_port)
        TIE_server.starttls()
        TIE_server.login(email_from,pswd)
        print("Succesfully connected to server")
        print()


    # Send emails to "person" as list is iterated
        print(f"Sending email to: {Name}...")
        TIE_server.sendmail(email_from, mail, text)
        print(f"Email sent to: {Name}")
        print()


        print(f'Correo número: {i+1}')
        print()

    # Close the port
        TIE_server.quit()

def extraer_subcadena(cadena, caracter_inicio, caracter_fin):
    subcadenas = []
    while caracter_inicio in cadena and caracter_fin in cadena:
        inicio = cadena.index(caracter_inicio) + 1
        fin = cadena.index(caracter_fin, inicio)
        subcadenas.append(cadena[inicio:fin])
        cadena = cadena[fin + 1:]
    return subcadenas

limit = context.shape[0]   

for i in range( 1 ):
    

    extract = context.loc[i]

    authors_email= ['edgaus98@gmail.com']
    Submission_Id = extract['Id']
        



    version = 'Datos\Sender Gmail Emails\Template.docx'
    doc = DocxTemplate(version)
    doc.render(extract)

    # Set the paths for the Word and PDF files
    word_path = f'Datos\Sender Gmail Emails\Documents_Generated\Word\Acceptance Letter-{Submission_Id}.docx'
    pdf_path = f'Datos\Sender Gmail Emails\Documents_Generated\PDF\Acceptance Letter-{Submission_Id}.pdf'
    pdf_name = f"Certificate Id-{Submission_Id}.pdf"


    doc.save(word_path)

    # Set the paths for the Word and PDF files
    word_path = f'Datos\Sender Gmail Emails\Documents_Generated\Word\Acceptance Letter-{Submission_Id}.docx'
    pdf_path = f'Datos\Sender Gmail Emails\Documents_Generated\PDF\Acceptance Letter-{Submission_Id}.pdf'
    pdf_name = f"Certificate Id-{Submission_Id}.pdf"
 
    # Load the Word document using the docx library
    doc = docx.Document(word_path)
 
    # Save the Word document as a PDF using Microsoft Word
    word = comtypes.client.CreateObject("Word.Application")
    docx_path = os.path.abspath(word_path)
    pdf_path = os.path.abspath(pdf_path)

 
    pdf_format = 17  # PDF file format code
    word.Visible = False
    in_file = word.Documents.Open(docx_path)
    in_file.SaveAs(pdf_path, FileFormat=pdf_format)
    in_file.Close()
  
    # Quit Microsoft Word
    word.Quit()


    version = 'Datos\Sender Gmail Emails\html.docx'
    version_SAVE = 'Datos\Sender Gmail Emails\html_SAVE.docx'

    doc = DocxTemplate(version)
    doc.render(extract)
    doc.save(version_SAVE)

    word.Quit()
    
    doc = docx.Document("Datos\Sender Gmail Emails\html_SAVE.docx")
    
    with open("file.txt", "w", encoding="utf-8") as f:
        for p in doc.paragraphs:
            f.write(p.text + "\n")

    with open("file.txt", "r", encoding="utf-8") as f:
        text = f.read()
    
    
    



    # Quit Microsoft Word
        
    
    # Run the function
    sender_html(extract,pdf_path,i,pdf_name,authors_email,'file.txt')

